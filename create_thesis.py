#!/usr/bin/env python3
"""
Generate Bachelor Thesis: Генерисање унарних SFPU кернела помоћу AI агената
Uses python-docx to create a formatted .docx from presentation content and narration.
"""

from docx import Document
from docx.shared import Pt, Inches, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ──────────────────────────────────────────────
# Constants
# ──────────────────────────────────────────────
IMG_DIR = "thesis_images"
OUTPUT = "Diplomski_rad.docx"

# Template page dimensions (EMU)
PAGE_WIDTH = 7560310
PAGE_HEIGHT = 10692130
MARGIN_LEFT = 900430
MARGIN_RIGHT = 899795
MARGIN_TOP = 1195705
MARGIN_BOTTOM = 899795


def setup_styles(doc):
    """Configure document styles to match ETF template."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(10)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.6)
    pf.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Cambria"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Caption style
    cap = doc.styles["Caption"]
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)


def setup_page(doc):
    """Set page dimensions, margins, headers and footers to match ETF template."""
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = Emu(449580)
    section.footer_distance = Emu(449580)

    # Enable different first page header/footer (title page has none)
    section.different_first_page_header_footer = True

    # -- First page header & footer: empty --
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    # Clear any default content
    for p in first_header.paragraphs:
        p.clear()

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    for p in first_footer.paragraphs:
        p.clear()

    # -- Default header (pages 2+): thesis title + horizontal rule --
    header = section.header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p_element = p._element
        p_element.getparent().remove(p_element)

    # Paragraph 1: thesis title
    hp = header.add_paragraph()
    hp.style = doc.styles["Header"]
    run = hp.add_run("Генерисање унарних SFPU кернела помоћу AI агената")
    run.font.size = Pt(10)

    # Paragraph 2: horizontal rule (VML-based, matching template)
    hr_p = header.add_paragraph()
    hr_p.style = doc.styles["Header"]
    hr_run = hr_p.add_run()
    # VML horizontal rule matching the template
    hr_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<v:rect id="_x0000_i1025" style="width:0;height:1.5pt" '
        'o:hralign="center" o:hrstd="t" o:hr="t" fillcolor="#a0a0a0" stroked="f"/>'
        "</w:pict>"
    )
    hr_run._element.append(parse_xml(hr_xml))

    # -- Default footer (pages 2+): centered page number --
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear default content
    for p in footer.paragraphs:
        p_element = p._element
        p_element.getparent().remove(p_element)

    # Empty line above page number
    fp1 = footer.add_paragraph()
    fp1.style = doc.styles["Footer"]

    # Centered page number field
    fp2 = footer.add_paragraph()
    fp2.style = doc.styles["Footer"]
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Build PAGE field: fldChar begin, instrText, fldChar separate, text, fldChar end
    _add_page_number_field(fp2)

    # Empty line below page number
    fp3 = footer.add_paragraph()
    fp3.style = doc.styles["Footer"]


def _add_page_number_field(paragraph):
    """Insert a PAGE number field into a paragraph, matching template structure."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # fldChar begin
    run1 = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1._element.append(fldChar_begin)

    # instrText
    run2 = paragraph.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">PAGE  </w:instrText>'
    )
    run2._element.append(instrText)

    # fldChar separate
    run3 = paragraph.add_run()
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_sep)

    # Placeholder text (Word updates this on open)
    run4 = paragraph.add_run("1")
    run4_rPr = run4._element.get_or_add_rPr()
    run4_rPr.append(parse_xml(f'<w:noProof {nsdecls("w")}/>'))

    # fldChar end
    run5 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar_end)


def add_page_break(doc):
    doc.add_page_break()


def add_centered_para(doc, text, font_name="Calibri", size=12, bold=False, space_before=0, space_after=0):
    """Add a centered paragraph with custom formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_normal_para(doc, text, bold=False, italic=False):
    """Add a normal justified paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_image_with_caption(doc, img_path, caption_text, width_inches=5.5):
    """Add an image centered with a caption below."""
    full_path = os.path.join(IMG_DIR, img_path)
    if not os.path.exists(full_path):
        p = doc.add_paragraph(f"[Слика недоступна: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(full_path, width=Inches(width_inches))

    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption_text)


def add_table(doc, headers, rows, caption_text=None):
    """Add a formatted table with optional caption above."""
    if caption_text:
        cap = doc.add_paragraph(style="Caption")
        cap.add_run(caption_text)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        # Shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    doc.add_paragraph()  # spacing
    return table


# ──────────────────────────────────────────────
# TITLE PAGE
# ──────────────────────────────────────────────
def _make_borderless_table(doc, rows, cols):
    """Create a borderless table matching the template style."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders via tblPr
    tblPr = tbl._tbl.tblPr
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders_xml)
    return tbl


def _no_spacing_centered(cell, text, size_pt=14, bold=False):
    """Write centered No Spacing text in a cell, matching the template."""
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles["No Spacing"] if hasattr(cell.part, "document") else p.style
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    return p


def write_title_page(doc):
    # ── University header table: [logo | university + katedra | empty] ──
    logo_path = os.path.join(IMG_DIR, "etf_logo.png")

    tbl = _make_borderless_table(doc, 1, 3)

    # Set column widths to match template (EMU)
    for cell, width in zip(tbl.rows[0].cells, [Emu(1383030), Emu(3703320), Emu(1411605)]):
        cell.width = width

    # Col 0: ETF logo
    cell0 = tbl.rows[0].cells[0]
    cell0.text = ""
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.first_line_indent = Pt(0)
    if os.path.exists(logo_path):
        run0 = p0.add_run()
        run0.add_picture(logo_path, width=Inches(1.0))

    # Col 1: University name + blank + Katedra
    cell1 = tbl.rows[0].cells[1]
    cell1.text = ""
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run("Универзитет у Београду - Електротехнички факултет")
    run1.font.size = Pt(14)

    # Blank line
    p1b = cell1.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1b.paragraph_format.first_line_indent = Pt(0)
    p1b.paragraph_format.space_before = Pt(0)
    p1b.paragraph_format.space_after = Pt(0)

    # Katedra
    p1c = cell1.add_paragraph()
    p1c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1c.paragraph_format.first_line_indent = Pt(0)
    p1c.paragraph_format.space_before = Pt(0)
    p1c.paragraph_format.space_after = Pt(0)
    run1c = p1c.add_run("Катедра за сигнале и системе")
    run1c.font.size = Pt(14)

    # Blank line
    p1d = cell1.add_paragraph()
    p1d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1d.paragraph_format.first_line_indent = Pt(0)

    # Col 2: empty
    cell2 = tbl.rows[0].cells[2]
    cell2.text = ""

    # ── Empty paragraphs before DIPLOMSKI RAD ──
    for _ in range(6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # ── ДИПЛОМСКИ РАД — 20pt bold centered ──
    add_centered_para(doc, "ДИПЛОМСКИ РАД", size=20, bold=True, space_after=4)

    # ── Thesis title — 20pt bold centered ──
    add_centered_para(
        doc,
        "Генерисање унарних SFPU кернела помоћу AI агената",
        size=20,
        bold=True,
        space_after=4,
    )

    # ── Empty paragraphs before candidate/mentor ──
    for _ in range(4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # ── Kandidat — centered, 14pt, bold label ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Кандидат")
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Владимир Игњатијевић, бр. индекса ЕР 2022/0006")
    run.font.size = Pt(14)

    # ── Blank line ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # ── Mentor — centered, 14pt, bold label ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Ментор")
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Марија Новичић, асистент")
    run.font.size = Pt(14)

    # ── Empty paragraphs before city/date ──
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # ── Beograd, mesec 2026. godine — 14pt centered ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("Београд, април 2026. године")
    run.font.size = Pt(14)


# ──────────────────────────────────────────────
# PREDGOVOR
# ──────────────────────────────────────────────
def write_predgovor(doc):
    add_page_break(doc)
    doc.add_heading("ПРЕДГОВОР", level=1)

    add_normal_para(
        doc,
        "Овај дипломски рад представља резултат практичног рада обављеног током стажирања "
        "у компанији Tenstorrent, на пројекту tt-metal. Рад је настао из потребе да се "
        "аутоматизује процес писања SFPU кернел кода за Tenstorrent AI акцелераторе, "
        "коришћењем AI агената базираних на великим језичким моделима.",
    )
    add_normal_para(
        doc,
        "Практични део дипломског рада је реализован уз менторство Славка Крстића "
        "из компаније Tenstorrent, који је надгледао развој agentic workflow-а и "
        "евалуацију резултата. Аутор се захваљује тиму за развој компајлера у компанији "
        "Tenstorrent на подршци и приступу хардверским ресурсима неопходним за тестирање.",
    )
    add_normal_para(
        doc,
        "Посебну захвалност дугујем менторки Марији Новичић са Електротехничког факултета "
        "Универзитета у Београду, на стручним саветима и вођењу кроз процес израде "
        "дипломског рада.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Владимир Игњатијевић")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("У Београду, април 2026.")
    run.font.size = Pt(12)


# ──────────────────────────────────────────────
# REZIME
# ──────────────────────────────────────────────
def write_rezime(doc):
    add_page_break(doc)
    doc.add_heading("РЕЗИМЕ РАДА", level=1)

    add_normal_para(
        doc,
        "Овај дипломски рад бави се аутоматизованим генерисањем SFPU (Special Floating Point Unit) "
        "кернел кода за Tenstorrent AI акцелераторе коришћењем AI агената. Tenstorrent чипови "
        "користе Tensix архитектуру са специјализованим SFPU јединицама намењеним за израчунавање "
        "нелинеарних и трансценденталних функција, попут активационих функција у неуронским мрежама. "
        "Иако се за сваку нову унарну операцију мора уредити дванаест фајлова, само један — SFPU "
        "кернел — садржи нову рачунску логику, што овај процес чини погодним за аутоматизацију.",
    )
    add_normal_para(
        doc,
        "У раду је развијен вишеагентни систем (agentic workflow) заснован на моделу Claude (Anthropic), "
        "који се састоји од оркестратора, проналазача референтних операција, анализатора, "
        "имплементатора, тестера и агента за саморефлексију. За побољшање поузданости агената "
        "имплементирани су hook механизми и breadcrumbs логовање. Систем је евалуиран на Kernel "
        "Bench платформи са 42 benchmark операције.",
    )
    add_normal_para(
        doc,
        "Резултати показују да је SFPU генератор успешно решио 4+1 нових операција (rpow, frac, "
        "swish, softshrink и softcap) и побољшао прецизност постојећих. Упоредна анализа са "
        "директним приступом (Raw Opus) показала је да је директан приступ 2,2 пута јефтинији "
        "и конзистентнији за појединачне операције, али да аgentic workflow пружа предности "
        "у структурираном приступу и вишеслојној оркестрацији.",
    )

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Кључне речи: ")
    run.bold = True
    p.add_run(
        "SFPU, AI агенти, генерисање кода, Tenstorrent, Tensix, agentic workflow, "
        "активационе функције, SFPI, кернел."
    )


# ──────────────────────────────────────────────
# SADRZAJ (placeholder - user generates in Word)
# ──────────────────────────────────────────────
def write_sadrzaj(doc):
    add_page_break(doc)
    doc.add_heading("САДРЖАЈ", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("[Садржај се генерише аутоматски у Word-у: References → Table of Contents]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


# ──────────────────────────────────────────────
# 1 УВОД
# ──────────────────────────────────────────────
def write_uvod(doc):
    add_page_break(doc)
    doc.add_heading("1 УВОД", level=1)

    add_normal_para(
        doc,
        "Развој специјализованих AI акцелератора представља један од кључних праваца "
        "у савременом рачунарству. Компанија Tenstorrent развија AI акцелераторе засноване "
        "на Tensix архитектури, која омогућава ефикасно извршавање тензорских операција "
        "кроз tile-базирано израчунавање на матрицама димензија 32×32. У оквиру ове "
        "архитектуре, Special Floating Point Unit (SFPU) представља специјализовану "
        "јединицу намењену за израчунавање нелинеарних и трансценденталних функција, "
        "укључујући активационе функције неуронских мрежа као што су sigmoid, tanh, "
        "ReLU и друге.",
    )
    add_normal_para(
        doc,
        "Писање SFPU кернел кода за нове унарне операције је репетитиван процес који "
        "захтева измену дванаест фајлова у кодној бази tt-metal, при чему само један "
        "фајл — сам SFPU кернел — садржи нову рачунску логику. Преосталих једанаест "
        "фајлова представљају регистрациони и инфраструктурни код који прати устаљене "
        "обрасце. Ова карактеристика чини процес погодним за аутоматизацију помоћу AI "
        "агената.",
    )
    add_normal_para(
        doc,
        "Циљ овог дипломског рада је развој и евалуација вишеагентног система (agentic "
        "workflow) за аутоматско генерисање SFPU кернел кода. Систем је заснован на "
        "великом језичком моделу Claude (Anthropic) и користи специјализоване агенте "
        "за анализу референтних операција, имплементацију кода и тестирање коректности. "
        "Поред самог генерисања кода, рад се бави и механизмима за управљање понашањем "
        "агената, укључујући hook механизме и breadcrumbs логовање.",
    )
    add_normal_para(
        doc,
        "Рад је организован на следећи начин. У поглављу 2 представљене су теоријске основе "
        "Tensix архитектуре, SFPU јединице и два приступа програмирању SFPU кернела (SFPI "
        "и TTI), као и методологија развоја agentic workflow-а. У поглављу 3 приказани су "
        "резултати евалуације на Kernel Bench платформи. У поглављу 4 дискутовани су "
        "добијени резултати и ограничења приступа, а у поглављу 5 изведени су закључци "
        "и дати правци будућег рада.",
    )


# ──────────────────────────────────────────────
# 2 ТЕОРИЈСКЕ ОСНОВЕ И МЕТОДОЛОГИЈА
# ──────────────────────────────────────────────
def write_metodologija(doc):
    add_page_break(doc)
    doc.add_heading("2 ТЕОРИЈСКЕ ОСНОВЕ И МЕТОДОЛОГИЈА", level=1)

    add_normal_para(
        doc,
        "У овом поглављу представљене су теоријске основе Tensix архитектуре и SFPU "
        "јединице, два приступа програмирању SFPU кернела, као и методологија развоја "
        "AI агената за аутоматско генерисање кода.",
    )

    # ── 2.1 Tensix ──
    doc.add_heading("2.1 Tensix архитектура", level=2)

    add_normal_para(
        doc,
        "Tensix је рачунска језгра Tenstorrent AI акцелератора у којој се извршавају "
        "све матричне операције, као и активационе функције. Свако Tensix језгро садржи "
        "пет RISC-V процесора: два за премештање података (dataflow), један за распакивање "
        "(unpack), један за математичке операције (math) и један за паковање (pack). Поред "
        "процесора, свако језгро садржи 1,5 MB SRAM меморије (L1), NoC интерфејсе за "
        "комуникацију, FPU (матричну јединицу) и SFPU (векторску јединицу) [1].",
    )
    add_normal_para(
        doc,
        "Израчунавање се заснива на tile-овима димензија 32×32 елемента. Подаци се из L1 "
        "меморије преносе у Dest регистар, где их SFPU обрађује, а резултати се враћају "
        "назад у Dest, а потом у L1. На слици 1 приказан је ток података кроз Tensix језгро.",
    )

    add_image_with_caption(
        doc,
        "slide3_Picture_2.png",
        "Слика 1. Ток података кроз Tensix језгро: од L1 меморије, "
        "кроз Unpacker, FPU/SFPU и Dest регистар, до Packer-а и назад у L1.",
        width_inches=5.8,
    )

    # ── 2.2 FPU и SFPU ──
    doc.add_heading("2.2 FPU и SFPU", level=2)

    add_normal_para(
        doc,
        "У оквиру Tensix језгра постоје две рачунске јединице: FPU (Floating Point Unit) и "
        "SFPU (Special Floating Point Unit). FPU је моћнија јединица у смислу throughput-а, "
        "намењена за линеарне операције као што су матрично множење (matmul) и редукције. "
        "Међутим, FPU не подржава нелинеарне, односно трансценденталне функције, нити "
        "рад са 32-битним форматима података [2].",
    )
    add_normal_para(
        doc,
        "С друге стране, SFPU може да обради мањи број података у јединици времена, али "
        "подржава 32-битне формате и користи се за израчунавање активационих функција "
        "попут sigmoid, tanh, exp и сличних. У табели 1 приказано је поређење ових "
        "двеју јединица.",
    )

    add_table(
        doc,
        ["Аспект", "FPU", "SFPU"],
        [
            ["Примена", "matmul, reduce", "recip, exp, sigmoid, tanh, cast"],
            ["32-битни формати", "НЕ", "ДА"],
            ["Програмирање", "TRISC thread 1", "SFPI C++ или TTI_* assembly"],
            ["Ограничења", "Без нелинеарних функција", "Нижи throughput"],
        ],
        caption_text="Табела 1. Поређење FPU и SFPU јединица.",
    )

    # ── 2.3 Архитектура SFPU-а ──
    doc.add_heading("2.3 Архитектура SFPU-а", level=2)

    add_normal_para(
        doc,
        "SFPU се састоји од осам инстанци, при чему је свака инстанца задужена за по "
        "две колоне Dest регистра. Свака инстанца располаже са 16 локалних регистара "
        "(LREG), од којих се у осам смештају подаци из Dest-а. Ових осам регистара се "
        "простире дуж четири lane-а, а сваки lane садржи по једну MAD (Multiply-Add) "
        "извршну јединицу. Тиме SFPU омогућава 32 паралелна израчунавања (8 инстанци × "
        "4 lane-а) [3].",
    )

    add_image_with_caption(
        doc,
        "slide6_Google_Shape;199;p45.png",
        "Слика 2. Структура SFPU-а: Dest регистар (16×16 бита, 1024 реда), "
        "8 SFPU инстанци са по 4 lane-а и 16 LREG регистара.",
        width_inches=4.5,
    )

    add_normal_para(
        doc,
        "Dest регистар представља централну структуру података за SFPU. Има димензије "
        "1024 × 16 × 16 бита и служи као извор улазних података и одредиште резултата "
        "SFPU израчунавања. Подаци се из Dest-а учитавају у LREG регистре (SFPLOAD), "
        "обрађују помоћу MAD извршних јединица, а затим враћају назад у Dest (SFPSTORE).",
    )

    add_image_with_caption(
        doc,
        "slide25_Picture_6.png",
        "Слика 3. Детаљна архитектура SFPU-а: Dest Slice-ови, SrcS Slice-ови, "
        "SFPU Slice-ови са Shared Front End-ом и Global Shift Network-ом.",
        width_inches=4.5,
    )

    # ── 2.4 Програмирање SFPU кернела ──
    doc.add_heading("2.4 Програмирање SFPU кернела", level=2)

    add_normal_para(
        doc,
        "Постоје два начина програмирања SFPU кернела: коришћењем TTI_* assembly "
        "инструкција (ниски ниво) и коришћењем SFPI компајлерског кода (виши ниво). "
        "SFPI (SFPU Programming Interface) је C++ апстракција која скрива детаље "
        "локалних регистара и адресних модова, чинећи код читљивијим и лакшим за "
        "одржавање [4].",
    )
    add_normal_para(
        doc,
        "На сликама 4 и 5 приказана је имплементација ReLU активационе функције у оба "
        "стила. Одмах се може приметити да је SFPI код знатно читљивији.",
    )

    add_image_with_caption(
        doc,
        "slide8_Picture_4.png",
        "Слика 4. Имплементација ReLU функције коришћењем TTI_* assembly инструкција.",
        width_inches=5.5,
    )

    add_image_with_caption(
        doc,
        "slide8_Picture_6.png",
        "Слика 5. Имплементација ReLU функције коришћењем SFPI компајлерског кода.",
        width_inches=4.5,
    )

    add_normal_para(
        doc,
        "TTI assembly приступ пружа две кључне могућности које SFPI не подржава. Прва су "
        "адресни модови (address modes), који омогућавају прецизно подешавање кретања "
        "показивача на Dest регистру из итерације у итерацију. Друга су Condition Code-ови "
        "(CC), механизам за условно извршавање дуж lane-ова. Међутим, CC код сложенијих "
        "функција знатно усложњава код и повећава вероватноћу грешке.",
    )
    add_normal_para(
        doc,
        "На сликама 6 и 7 приказан је пример clamping функције, чија имплементација "
        "јасно показује разлику у сложености између TTI и SFPI приступа.",
    )

    add_image_with_caption(
        doc,
        "slide9_Picture_4.png",
        "Слика 6. Clamping функција у TTI_* assembly — сложена употреба Condition Code-ова.",
        width_inches=5.5,
    )

    add_image_with_caption(
        doc,
        "slide9_Picture_10.png",
        "Слика 7. Clamping функција у SFPI — једноставна употреба v_if/v_elseif/v_endif.",
        width_inches=4.2,
    )

    add_normal_para(doc, "У табели 2 дато је систематско поређење два приступа програмирању SFPU кернела.")

    add_table(
        doc,
        ["Приступ", "Предности", "Недостаци", "Када користити"],
        [
            [
                "SFPI\n(компајлер)",
                "Читљив C++ код (vFloat, v_if)\nПортабилан између HW генерација\nv_if / v_endif",
                "Без uint16/uint8 подршке\nБез cross-lane операција\nНе може директно Dest",
                "Подразумевано за нове кернеле — читљивост и коректност",
            ],
            [
                "TTI_*\n(assembly)",
                "Cross-lane операције\nЕксплицитна Dest манипулација\nВишестепени CC прелази",
                "Тешко читљив/одржив\nНије портабилан\nCC push/pop склон грешкама",
                "Када треба: uint16, SFPTRANSP, SFPSWAP, или максималан throughput",
            ],
        ],
        caption_text="Табела 2. Поређење SFPI и TTI приступа програмирању SFPU кернела.",
    )

    add_normal_para(
        doc,
        "На основу анализе, SFPI код треба користити кадгод је то могуће, због "
        "читљивости и преносивости. TTI assembly треба користити само у следећим "
        "случајевима: (1) када су потребне изузетно високе перформансе, (2) када треба "
        "манипулисати SFPU на нивоу lane-ова и регистара (SFPTRANSP, SFPSWAP), (3) "
        "када се користе адресни модови, и (4) за обраду uint16 и мањих формата [4].",
    )

    # ── 2.5 AI агенти ──
    doc.add_heading("2.5 AI агенти", level=2)

    add_normal_para(
        doc,
        "AI агенти, у контексту овог рада, представљају инстанце великог језичког модела "
        "Claude (Anthropic), које су током пројектовања промптоване инструкцијама записаним "
        "у одговарајућем markdown фајлу, заједно са tool-овима (алатима) који су им потребни "
        "за извршавање задатих задатака. Агент се разликује од обичног ChatBot-а по томе "
        "што има приступ external tool-овима и може аутономно да извршава акције у "
        "софтверском окружењу [5].",
    )
    add_normal_para(
        doc,
        "Рад агената може се побољшати на више начина. Први је додавање tool-ова, "
        "посебно MCP (Model Context Protocol) tool-ова, попут DeepWiki, Glean и "
        "Atlassian конектора. Ови tool-ови омогућавају агенту приступ документацији "
        "о хардверу и интерним ресурсима компаније који нису јавно доступни.",
    )
    add_normal_para(
        doc,
        "Други начин је коришћење подагената (subagent mechanism). Подзадаци се могу "
        "доделити подагентима, који своје крајње резултате прослеђују родитељском агенту. "
        "Овим се постиже боље управљање контекстом и паралелизација рада.",
    )

    # ── 2.6 Механизми за управљање агентима ──
    doc.add_heading("2.6 Механизми за управљање агентима", level=2)

    add_normal_para(
        doc,
        "Током развоја agentic workflow-а идентификована су три кључна проблема: "
        "(1) оркестратор не сачека резултате својих подагената већ одради посао уместо "
        "њих, (2) подагенти забораве да пишу излазне фајлове, и (3) тешко је пратити "
        "рад самих подагената. За решавање ових проблема развијена су два механизма: "
        "hook-ови и breadcrumbs.",
    )

    # Hooks
    p = doc.add_paragraph()
    run = p.add_run("Hook-ови")
    run.bold = True
    p.add_run(
        " су корисничке скрипте које се активирају на дефинисане догађаје из животног "
        "циклуса агента. У зависности од дефиниције, могу се активирати пре или после "
        "коришћења tool-ова, на почетку или крају агентовог рада, или пре/после "
        "компактовања контекста. Hook-ови се додају у оквиру YAML описа у заглављу "
        "агента. На слици 8 приказан је пример hook-а који блокира завршетак рада "
        "анализатора уколико постоје некомитоване промене."
    )

    add_image_with_caption(
        doc,
        "slide17_Picture_12.png",
        "Слика 8. Пример hook дефиниције у YAML заглављу агента: "
        "block_if_uncommitted скрипта се покреће при завршетку рада анализатора.",
        width_inches=5.5,
    )

    # Breadcrumbs
    p = doc.add_paragraph()
    run = p.add_run("Breadcrumbs")
    run.bold = True
    p.add_run(
        " су logging систем развијен у оквиру тима, који обавезује агенте да исписују "
        "све битне догађаје у JSONL формату. Систем се састоји од две скрипте: "
        "init_breadcrumbs.sh за иницијализацију JSONL фајла и append_breadcrumbs.sh "
        "за додавање догађаја. Breadcrumbs омогућавају праћење агентовог понашања кроз "
        "време — шта га успорава, које проблеме наилази, и колико времена му треба "
        "за сваки корак."
    )

    add_image_with_caption(
        doc,
        "slide18_Picture_21.png",
        "Слика 9. Документација breadcrumbs система: иницијализација и "
        "правила логовања за TTNN SFPU Operation Tester агента.",
        width_inches=5.5,
    )

    add_image_with_caption(
        doc,
        "slide19_Picture_13.png",
        "Слика 10. Пример breadcrumbs лога: JSONL записи са хипотезама, "
        "применама поправки и статусима тестирања за softsign операцију.",
        width_inches=5.8,
    )

    # ── 2.7 Agentic workflow ──
    doc.add_heading("2.7 Agentic workflow за генерисање SFPU кернела", level=2)

    add_normal_para(
        doc,
        "Развијени agentic workflow за генерисање SFPU кернела састоји се од следећих "
        "компоненти, организованих у хијерархијску структуру:",
    )

    # Workflow components as list
    components = [
        (
            "Оркестратор (Generator)",
            "Главни агент који координира целокупан процес. "
            "Прима назив нове операције и математичку дефиницију, затим покреће остале агенте "
            "по дефинисаном редоследу.",
        ),
        (
            "Проналазач референци (Reference Discoverer)",
            "Анализира математичку дефиницију "
            "нове операције и проналази 5 најрелевантнијих постојећих операција у кодној бази "
            "које могу послужити као референца.",
        ),
        (
            "Анализатори (5× Analyzer)",
            "За сваку референтну операцију покреће се по један "
            "анализатор који детаљно проучава имплементацију SFPU кернела и генерише анализу.",
        ),
        (
            "Имплементатор (Implementor)",
            "На основу анализа, имплементира нову операцију "
            "кроз свих 12 апстракцијских слојева, укључујући SFPU кернел.",
        ),
        (
            "Тестер (Tester)",
            "Креира тестове за нову операцију, покреће их и итерира " "на грешкама док операција не прође све тестове.",
        ),
        (
            "Агент за саморефлексију (Self-Reflection)",
            "Анализира целокупан процес, " "идентификује проблеме и предлаже побољшања за будуће покретања.",
        ),
    ]

    for name, desc in components:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f"• {name} — ")
        run.bold = True
        p.add_run(desc)

    add_normal_para(
        doc,
        "Једна од кључних одлука коју агенти доносе јесте избор између SFPI и TTI "
        "приступа за имплементацију SFPU кернела. За већину унарних операција, SFPI "
        "је препоручен избор због читљивости и коректности.",
    )


# ──────────────────────────────────────────────
# 3 РЕЗУЛТАТИ
# ──────────────────────────────────────────────
def write_rezultati(doc):
    add_page_break(doc)
    doc.add_heading("3 РЕЗУЛТАТИ", level=1)

    add_normal_para(
        doc,
        "У овом поглављу приказани су резултати евалуације развијеног agentic workflow-а "
        "на Kernel Bench платформи. Прво је дат пример операције RReLU, а затим "
        "су представљени збирни резултати прецизности и упоредна анализа.",
    )

    # ── 3.1 RReLU ──
    doc.add_heading("3.1 Пример: RReLU операција", level=2)

    add_normal_para(
        doc,
        "RReLU (Randomized Leaky Rectified Linear Unit) представља стохастичку варијанту "
        "Leaky ReLU активационе функције. Дефинисана је на следећи начин:",
    )

    add_image_with_caption(
        doc,
        "slide21_Picture_12.png",
        "Слика 11. Математичка дефиниција RReLU функције: за x ≥ 0 враћа x, "
        "а за x < 0 враћа ax, где је a параметар нагиба.",
        width_inches=3.5,
    )

    add_normal_para(
        doc,
        "Током тренирања, параметар a се бира случајно из униформне расподеле на интервалу "
        "[lower, upper], док се током евалуације користи детерминистичка вредност "
        "a = (lower + upper) / 2. На слици 12 приказан је график RReLU функције "
        "са параметрима lower=1/8 и upper=1/3.",
    )

    add_image_with_caption(
        doc,
        "slide21_Picture_10.png",
        "Слика 12. График RReLU функције (lower=1/8, upper=1/3): црвена линија "
        "за x ≥ 0, љубичаста зона представља опсег стохастичког нагиба.",
        width_inches=4.0,
    )

    # ── 3.2 Kernel Bench ──
    doc.add_heading("3.2 Kernel Bench платформа", level=2)

    add_normal_para(
        doc,
        "Kernel Bench је конкурентна benchmark платформа за евалуацију AI-генерисаних "
        "кернела на Tenstorrent хардверу. Евалуација је спроведена на N300 хардверској "
        "платформи са укупно 42 benchmark операције подељене у шест категорија.",
    )

    add_image_with_caption(
        doc,
        "slide22_Picture_7.png",
        "Слика 13. Расподела 42 benchmark операције по статусу: решене (3), "
        "у току (15), неуспешне (20), ненападнуте (4).",
        width_inches=3.5,
    )

    add_image_with_caption(
        doc,
        "slide22_Picture_8.png",
        "Слика 14. Расподела benchmark операција по категоријама: унарне активације (10), "
        "унарна математика (11), редукције (11), померање података (8), "
        "gated јединице (4), мулти-операнд (1).",
        width_inches=4.5,
    )

    add_normal_para(
        doc,
        "Међу најуспешнијим операцијама су: hardsigmoid, hardswish и softsign са 100% "
        "pass rate, hardtanh са 97,7%, и cosh и sinh са 68,7%. Ове операције "
        "користе релативно једноставне математичке функције погодне за SFPI имплементацију. "
        "Прецизност представља примарну баријеру за решавање преосталих операција.",
    )

    # ── 3.3 Прецизност ──
    doc.add_heading("3.3 Резултати прецизности — SFPU Generator vs Main", level=2)

    add_normal_para(
        doc,
        "У табели 3 приказани су резултати упоређивања SFPU Generator-а са основном "
        "граном (main) tt-metal репозиторијума за девет одабраних операција. Резултати "
        "показују да је SFPU генератор успешно решио четири нове операције (rpow, frac, "
        "swish, softshrink) плус softcap као нови benchmark, и побољшао cbrt са 0% на 82,1%.",
    )

    add_table(
        doc,
        ["Операција", "Main Leaderboard", "metal_main Pass Rate", "sfpu-generator Pass Rate", "Статус"],
        [
            ["rpow", "52,5%", "644/644 (100,0%)", "644/644 (100,0%)", "НОВО решено"],
            ["frac", "63,1%", "271/271 (100,0%)", "271/271 (100,0%)", "НОВО решено"],
            ["swish", "58,3%", "5/24 (20,8%)", "24/24 (100,0%)", "НОВО решено"],
            ["softshrink", "41,0%", "0/1856 (0,0%)", "1856/1856 (100,0%)", "НОВО решено"],
            ["softcap", "Н/А (нови)", "Н/А", "203/203 (100,0%)", "НОВО решено"],
            ["hardsigmoid", "100,0%", "195/261 (74,7%)", "261/261 (100,0%)", "Већ решено"],
            ["hardswish", "100,0%", "195/259 (75,3%)", "259/259 (100,0%)", "Већ решено"],
            ["softsign", "100,0%", "229/229 (100,0%)", "229/229 (100,0%)", "Већ решено"],
            ["cbrt", "0%", "59/273 (21,6%)", "224/273 (82,1%)", "Побољшано"],
        ],
        caption_text="Табела 3. Поређење SFPU Generator-а и main гране за одабране операције.",
    )

    add_normal_para(
        doc,
        "Четири ново решене операције (rpow, frac, swish, softshrink) постижу 100% pass "
        "rate, што значи да све тест инстанце пролазе проверу прецизности. Softcap "
        "операција је такође решена са 100% pass rate, али је представљена као нови "
        "benchmark који није постојао на main грани. За операције које су већ биле "
        "решене (hardsigmoid, hardswish, softsign), SFPU генератор постиже исте или "
        "боље резултате.",
    )

    # ── 3.4 Softcap евалуација ──
    doc.add_heading("3.4 Евалуација softcap операције", level=2)

    add_normal_para(
        doc,
        "За детаљнију евалуацију спроведено је поређење два приступа на softcap операцији: "
        "директан приступ са Raw Opus моделом (без специјализованих агената) и SFPU "
        "Generator workflow. Оба приступа су покренута по 5 пута, са приступом "
        "дестилованом знању из кодне базе, без приступа интернету, и евалуирана "
        "на истим golden тестовима.",
    )

    add_image_with_caption(
        doc,
        "slide24_Picture_21.png",
        "Слика 15. Агрегатно поређење Raw Opus и SFPU Generator-а (по 5 покретања): "
        "просечан pass rate, просечна цена и укупан број пролазних golden тестова.",
        width_inches=5.0,
    )

    add_normal_para(
        doc,
        "Резултати показују да Raw Opus остварује 4/5 савршених покретања у поређењу "
        "са 2/5 за SFPU Generator. Просечна цена по покретању износи $29,7 за Raw Opus "
        "и $66,8 за SFPU Generator, што чини директан приступ 2,2 пута јефтинијим. "
        "Raw Opus такође показује мањи tail risk, односно конзистентније резултате.",
    )


# ──────────────────────────────────────────────
# 4 ДИСКУСИЈА
# ──────────────────────────────────────────────
def write_diskusija(doc):
    add_page_break(doc)
    doc.add_heading("4 ДИСКУСИЈА", level=1)

    add_normal_para(
        doc,
        "Резултати приказани у претходном поглављу показују да је аутоматско генерисање "
        "SFPU кернел кода помоћу AI агената остварив приступ, са значајним успесима "
        "у решавању нових операција. Међутим, упоредна анализа на softcap операцији "
        "открива занимљиве trade-off-ове између вишеагентног и директног приступа.",
    )
    add_normal_para(
        doc,
        "SFPU Generator, као вишеагентни систем, пружа структуриран приступ проблему: "
        "аутоматски проналази релевантне референтне операције, анализира их, и на основу "
        "анализа генерише нову имплементацију. Овај приступ је посебно користан за "
        "инжењере који нису детаљно упознати са SFPU архитектуром, јер агенти сами "
        "доносе одлуке о избору приступа (SFPI или TTI) и структури кода.",
    )
    add_normal_para(
        doc,
        "С друге стране, директан приступ (Raw Opus) показао се ефикаснијим на појединачним "
        "операцијама — 2,2 пута јефтинији и конзистентнији. Ово указује на то да додатна "
        "сложеност вишеагентног система уноси overhead који се не оправдава увек, "
        "посебно за операције где је математичка дефиниција јасна и једноставна.",
    )
    add_normal_para(
        doc,
        "Прецизност представља примарну баријеру за решавање преосталих операција на "
        "Kernel Bench платформи. Многе нерешене операције захтевају полиномске "
        "апроксимације трансценденталних функција, где мале грешке у коефицијентима "
        "могу довести до значајних одступања. Breadcrumbs логовање показало се кључним "
        "за дијагностику ових проблема, омогућавајући праћење хипотеза и поправки "
        "које агенти примењују (видети слику 10).",
    )
    add_normal_para(
        doc,
        "Hook механизми решили су два значајна проблема: спречавање превременог завршетка "
        "оркестратора без чекања подагената, и осигуравање да подагенти запишу своје "
        "излазне фајлове. Без ових механизама, workflow је био непоуздан и тешко "
        "репродуцибилан.",
    )
    add_normal_para(
        doc,
        "Ограничења овог рада укључују: (1) евалуација је спроведена само на унарним "
        "eltwise операцијама, (2) софтcap евалуација је базирана на релативно малом "
        "узорку од 5 покретања по приступу, и (3) цена покретања зависи од тренутних "
        "цена API позива великих језичких модела.",
    )


# ──────────────────────────────────────────────
# 5 ЗАКЉУЧАК
# ──────────────────────────────────────────────
def write_zakljucak(doc):
    add_page_break(doc)
    doc.add_heading("5 ЗАКЉУЧАК", level=1)

    add_normal_para(
        doc,
        "У овом дипломском раду развијен је и евалуиран вишеагентни систем за "
        "аутоматско генерисање SFPU кернел кода за Tenstorrent AI акцелераторе. "
        "Систем користи специјализоване AI агенте базиране на моделу Claude за "
        "анализу референтних операција, имплементацију нових кернела и тестирање "
        "коректности.",
    )
    add_normal_para(
        doc,
        "Главни доприноси рада су: (1) успешно генерисање 4+1 нових SFPU операција "
        "(rpow, frac, swish, softshrink и softcap) које нису биле решене на основној "
        "грани tt-metal репозиторијума, (2) побољшање прецизности постојећих операција "
        "(cbrt са 0% на 82,1%), (3) развој hook механизама и breadcrumbs логовања "
        "као кључних компоненти за поузданост вишеагентних система, и (4) упоредна "
        "анализа вишеагентног и директног приступа генерисању кода.",
    )
    add_normal_para(
        doc,
        "Упоредна анализа показала је да директан приступ (Raw Opus) надмашује "
        "вишеагентни систем у погледу цене (2,2× јефтинији) и конзистентности "
        "за појединачне операције. Међутим, вишеагентни систем пружа предности "
        "у структурираном приступу и аутоматизацији целокупног процеса, укључујући "
        "проналажење референци и вишеслојну оркестрацију.",
    )
    add_normal_para(
        doc,
        "Правци будућег рада укључују: (1) поједностављивање agentic flow-а ради "
        "смањења overhead-а и цене, (2) проширивање документације о полиномским "
        "апроксимацијама како би се побољшала прецизност генерисаних кернела, (3) "
        "примену приступа на сложеније категорије операција (редукције, померање "
        "података), и (4) интеграцију директног и вишеагентног приступа у хибридни "
        "систем који бира стратегију на основу сложености операције.",
    )


# ──────────────────────────────────────────────
# 6 ЛИТЕРАТУРА
# ──────────────────────────────────────────────
def write_literatura(doc):
    add_page_break(doc)
    doc.add_heading("6 ЛИТЕРАТУРА", level=1)

    references = [
        'Tenstorrent Inc., "TT-Metal: Open-source AI accelerator programming framework," '
        "GitHub repository, 2024. [Online]. Available: https://github.com/tenstorrent/tt-metal.",
        'Tenstorrent Inc., "Metalium Guide," техничка документација у оквиру tt-metal '
        "репозиторијума, METALIUM_GUIDE.md, 2024.",
        'Tenstorrent Inc., "SFPU ISA Documentation," интерна Confluence документација, 2024.',
        'Tenstorrent Inc., "SFPI Programming Interface," изворни код у tt-metal/tt_metal/hw/, 2024.',
        'Anthropic, "Claude: AI Assistant," 2024. [Online]. Available: https://www.anthropic.com/claude.',
        'Anthropic, "Model Context Protocol (MCP)," техничка спецификација, 2024. '
        "[Online]. Available: https://modelcontextprotocol.io/.",
        'A. Paszke et al., "PyTorch: An Imperative Style, High-Performance Deep Learning Library," '
        "in Advances in Neural Information Processing Systems, vol. 32, 2019, pp. 8024-8035.",
        'K. Xu et al., "Empirical Evaluation of Rectified Activations in Convolutional Network," '
        "arXiv preprint arXiv:1505.00853, 2015.",
        'M. Chen et al., "Evaluating Large Language Models Trained on Code," arXiv preprint ' "arXiv:2107.03374, 2021.",
        'Y. Li et al., "Competition-Level Code Generation with AlphaCode," Science, vol. 378, '
        "no. 6624, pp. 1092-1097, Dec. 2022.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.hanging_indent = Cm(1.0)
        run = p.add_run(f"[{i}] ")
        run.bold = True
        p.add_run(ref)


# ──────────────────────────────────────────────
# ПРИЛОГ А
# ──────────────────────────────────────────────
def write_prilog(doc):
    add_page_break(doc)
    doc.add_heading("ПРИЛОГ А", level=1)

    add_normal_para(
        doc, "У овом прилогу дати су додатни примери кода и детаљни дијаграми архитектуре " "SFPU јединице."
    )

    doc.add_heading("А.1 Dest регистар", level=2)

    add_image_with_caption(
        doc, "slide27_Picture_4.png", "Слика А1. Dest Regfile: 1024 реда × 16 колона × 16 бита.", width_inches=3.5
    )

    doc.add_heading("А.2 Адресни модови SFPU-а", level=2)

    add_image_with_caption(
        doc,
        "slide27_Picture_45.png",
        "Слика А2. SFPU адресни модови — детаљан приказ итерирања кроз Dest регистар.",
        width_inches=5.0,
    )

    doc.add_heading("А.3 Детаљна архитектура SFPU инстанце", level=2)

    add_image_with_caption(
        doc,
        "slide14_Picture_4.png",
        "Слика А3. Хардверски модел SFPU-а из ISA документације: " "SFPU Slice-ови, SrcS банке и MAD извршне јединице.",
        width_inches=4.5,
    )

    add_image_with_caption(
        doc,
        "slide14_Picture_3.png",
        "Слика А4. Додатни хардверски дијаграм из SFPU ISA документације.",
        width_inches=4.0,
    )


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────
def main():
    doc = Document()

    setup_styles(doc)
    setup_page(doc)

    write_title_page(doc)
    write_predgovor(doc)
    write_rezime(doc)
    write_sadrzaj(doc)
    write_uvod(doc)
    write_metodologija(doc)
    write_rezultati(doc)
    write_diskusija(doc)
    write_zakljucak(doc)
    write_literatura(doc)
    write_prilog(doc)

    doc.save(OUTPUT)
    print(f"Дипломски рад сачуван: {OUTPUT}")


if __name__ == "__main__":
    main()
